#!/usr/bin/env python3
"""Convert all .docx files in docs/ to Markdown (Chinese and English versions)."""

import os
import re
import docx

DOCS_DIR = "docs"
ZH_DIR = os.path.join(DOCS_DIR, "zh")
EN_DIR = os.path.join(DOCS_DIR, "en")

# ── Running headers to strip from English file ──
EN_RUNNING_HEADERS = {
    # Part headers
    "Building with God",
    "The Spirit- Guided Parent",
    "The Spirit-Guided Parent",
    "Spirit- Filled Children",
    "Spirit-Filled Children",
    "Conclusion",
    # Chapter headers
    "Foreword",
    "Introduction",
    "The House That God Builds",
    "The Way They Should Go",
    "Stay a Novice",
    "Raising Sons and Daughters",
    "Dreaming with God",
    "Do Not (Always) Tell the Truth",
    "Learning God's Alphabet",
    "Children Who See",
    "Exercising Our Senses",
    "Word Becoming Flesh",
    "They Have Been with Jesus",
    "Staying Hungry",
    "Fighting and Building",
}

EN_SKIP_PATTERNS = [
    "(Unpublished manuscript—copyright protected Baker Publishing Group)",
]

# ── English chapter boundaries (paragraph indices) ──
# Format: (start_idx, chapter_id, title, part_info, sections)
EN_CHAPTERS = [
    # Endorsements + Dedication (cover)
    (0, "00-cover-and-dedication", "Cover and Dedication", None, []),
    # Foreword
    (21, "00-foreword", "Foreword", None, []),
    # Introduction
    (34, "00-introduction", "Introduction", None, []),
    # Part 1: Building with God
    (69, "01-the-house-that-god-builds", "The House That God Builds",
     "Part 1: Building with God",
     ["Creating Culture", "Education through Demonstration",
      "The Need for Experience", "Reflecting on the Experiences",
      "Parenting Like the Father", "Prayer Prompt"]),
    (141, "02-the-way-they-should-go", "The Way They Should Go", None,
     ["The Importance of Encounter", "Training Is More Than Explaining",
      "What We Allow In", "The Parenting Prayer"]),
    (188, "03-stay-a-novice", "Stay a Novice", None,
     ["Afraid to Father a Son", "Buy That Boy Some Throwing Knives",
      "Building Trust", "Ask the Father", "Beach House for a Month", "The Novice"]),
    # Part 2: The Spirit-Guided Parent
    (266, "04-raising-sons-and-daughters", "Raising Sons and Daughters",
     "Part 2: The Spirit-Guided Parent",
     ["The Orphanage in the Mind"]),
    (342, "05-dreaming-with-god", "Dreaming with God", None,
     ["Breaking Cycles", "Downloads Available", "Swapping Magazines",
      "Replacing the Magazine"]),
    (401, "06-do-not-always-tell-the-truth", "Do Not (Always) Tell the Truth", None,
     ["Question the Lies", "I Have Thought That, Too", "Let's Flush It",
      "Destroying Enemies", "But I Have Prayed"]),
    # Part 3: Spirit-Filled Children
    (493, "07-learning-gods-alphabet", "Learning God's Alphabet",
     "Part 3: Spirit-Filled Children",
     ["Asking More Questions", "How God Talks", "God Speaks in Our Hearts",
      "How Many Children Can Prophesy?", "Prayer Prompts"]),
    (571, "08-when-children-see", "When Children See", None,
     ["Keep the Door Shut", "We Want Our Kids to See", "Again with the Questions",
      "Seeing Is a Gift That Needs to Be Developed"]),
    (636, "09-training-our-senses", "Training Our Senses", None,
     ["The Word of Righteousness", "Exercise Is Profitable", "Feeling and Hearing",
      "Trained or Tormented", "Rise of the Guardians",
      "Maintaining Our Atmosphere", "Angels Everywhere", "Heaven at Home"]),
    (728, "10-the-word-became-flesh", "The Word Became Flesh", None,
     ["Into His Word", "Let the Little Children Come"]),
    (767, "11-they-have-been-with-jesus", "They Have Been with Jesus", None,
     ["Like a Boss", "Lesson Planning", "Swallowing Up Cancer",
      "Do What You See Him Doing", "Prayer Prompts"]),
    # Part 4: Conclusion
    (822, "12-staying-hungry", "Staying Hungry",
     "Part 4: Conclusion",
     ["The Grace of God", "How I Got What I Have", "Stir Up What You Have"]),
    (854, "13-fighting-and-building", "Fighting and Building", None,
     ["A Table Is Prepared", "It Is Time!"]),
]

# ── Chinese file mapping ──
ZH_FILES = [
    ("00(Ada)_封面封底-3光哥改.docx", "00-封面與獻詞", "封面與獻詞", None),
    ("00(唐寧)_引言-2光哥改.docx", "00-序與引言", "序與引言", None),
    ("01(唐寧)_神所建造的殿(光哥改).docx", "01-神所建造的殿", "神所建造的殿", "第一部：與神一同建造"),
    ("02(唐寧)_孩子們所當行的路(光哥改).docx", "02-孩子們所當行的路", "孩子們所當行的路", None),
    ("03(唐寧)_保有初心(光哥改).docx", "03-保有初心", "保有初心", None),
    ("04(張聖晞)_興起兒女(光哥改).docx", "04-興起兒女", "興起兒女", "第二部：跟隨聖靈引導的父母"),
    ("05(張聖晞)_與神一同築夢(光哥改).docx", "05-與神一同築夢", "與神一同築夢", None),
    ("06(張聖晞)_不要（一直）講真理.(Sarah潤)docx(光哥改).docx", "06-不要一直講真理", "不要（一直）講真理", None),
    ("07(張聖晞)_學習神怎麼拼字.docx.docx", "07-學習神怎麼拼字", "學習神怎麼拼字", "第三部：被聖靈充滿的孩子"),
    ("08(Ada)_當孩子們看見(光哥改).docx", "08-當孩子們看見", "當孩子們看見", None),
    ("09(Ada)_操練我們的感官(光哥改).docx", "09-操練我們的感官", "操練我們的感官", None),
    ("10(唐寧)_道成肉身(光哥改).docx", "10-道成肉身", "道成肉身", None),
    ("11(Ada)_他們有跟過耶穌_new(光哥改).docx", "11-他們有跟過耶穌", "他們有跟過耶穌", None),
    ("12(Ada)_保持渴慕(光哥改).docx", "12-保持渴慕", "保持渴慕", "第四部：結論"),
    ("13(Ada)_爭戰與建造(光哥改).docx", "13-爭戰與建造", "爭戰與建造", None),
]

# Chinese section titles to detect as ## headings
ZH_KNOWN_SECTIONS = {
    # Ch1
    "真實的文化", "透過示範來教導", "經歷的必要性", "回顧過去的經歷", "像天父一樣做父母",
    # Ch2
    "經歷的重要性", "訓練不只是解釋", "我們允許什麼進來", "父母的禱告",
    # Ch3
    "我們並不知道自己在做什麼", "害怕為人父", "買飛刀給那個男孩", "建立信任",
    "問天父", "海邊小屋一個月", "新手",
    # Ch4
    "腦海中的孤兒院",
    # Ch5
    "打破循環", "可下載的內容", "交換彈匣", "更換彈匣",
    # Ch6
    "質疑謊言", "我也曾這樣想過", "沖掉它吧", "消滅仇敵", "但我已經禱告了",
    # Ch7
    "問更多問題", "神如何說話", "神在我們心裡說話", "有多少孩子可以說預言？",
    # Ch8
    "把門關上", "我們希望孩子們能看見", "再來問問題", "看見是一個需要被發展的恩賜",
    # Ch9
    "公義的話語", "操練是有益的", "感覺和聽見", "受訓練還是受折磨",
    "乃是靠著信心和忍耐", "守護者聯盟", "維護我們的氛圍", "到處都是天使", "家中的天堂",
    # Ch10
    "進入祂的話語", "讓小孩子到我這裡來",
    # Ch11
    "像老闆一樣", "課程規劃", "吞噬癌症", "做你看見祂所做的",
    # Ch12
    "神的恩典", "我怎麼得到我所擁有的", "攪動你所擁有的",
    # Ch13
    "筵席已經預備好了", "是時候了！",
}

# ── Patterns for detecting scripture references (for blockquote) ──
SCRIPTURE_PATTERN = re.compile(
    r'^[「『]?.{0,5}(創世記|出埃及記|利未記|民數記|申命記|約書亞記|士師記|路得記|'
    r'撒母耳記|列王紀|歷代志|以斯拉記|尼希米記|以斯帖記|約伯記|詩篇|箴言|傳道書|'
    r'雅歌|以賽亞書|耶利米書|耶利米哀歌|以西結書|但以理書|何西阿書|約珥書|阿摩司書|'
    r'俄巴底亞書|約拿書|彌迦書|那鴻書|哈巴谷書|西番雅書|哈該書|撒迦利亞書|瑪拉基書|'
    r'馬太福音|馬可福音|路加福音|約翰福音|使徒行傳|羅馬書|哥林多前書|哥林多後書|'
    r'加拉太書|以弗所書|腓立比書|歌羅西書|帖撒羅尼迦前書|帖撒羅尼迦後書|'
    r'提摩太前書|提摩太後書|提多書|腓利門書|希伯來書|雅各書|彼得前書|彼得後書|'
    r'約翰一書|約翰二書|約翰三書|猶大書|啟示錄)',
    re.UNICODE
)

EN_SCRIPTURE_PATTERN = re.compile(
    r'^(Genesis|Exodus|Leviticus|Numbers|Deuteronomy|Joshua|Judges|Ruth|'
    r'1\s*Samuel|2\s*Samuel|1\s*Kings|2\s*Kings|1\s*Chronicles|2\s*Chronicles|'
    r'Ezra|Nehemiah|Esther|Job|Psalm|Psalms|Proverbs|Ecclesiastes|'
    r'Song\s*of\s*Solomon|Isaiah|Jeremiah|Lamentations|Ezekiel|Daniel|Hosea|Joel|Amos|'
    r'Obadiah|Jonah|Micah|Nahum|Habakkuk|Zephaniah|Haggai|Zechariah|Malachi|'
    r'Matthew|Mark|Luke|John|Acts|Romans|1\s*Corinthians|2\s*Corinthians|'
    r'Galatians|Ephesians|Philippians|Colossians|1\s*Thessalonians|2\s*Thessalonians|'
    r'1\s*Timothy|2\s*Timothy|Titus|Philemon|Hebrews|James|1\s*Peter|2\s*Peter|'
    r'1\s*John|2\s*John|3\s*John|Jude|Revelation)\s+\d',
    re.IGNORECASE
)


def is_zh_section_title(text, prev_empty, next_empty):
    """Detect if a paragraph is a section title in Chinese content."""
    stripped = text.strip()
    if not stripped:
        return False
    # Only use known section titles - heuristic detection causes false positives on dialogue
    if stripped in ZH_KNOWN_SECTIONS:
        return True
    return False


def is_scripture_ref_zh(text):
    """Check if line is a standalone scripture reference."""
    stripped = text.strip()
    if SCRIPTURE_PATTERN.search(stripped):
        # Only if it's relatively short (a reference, not a full paragraph)
        if len(stripped) < 60:
            return True
    return False


def convert_zh_file(filepath, title, part_info):
    """Convert a Chinese .docx file to Markdown string."""
    doc = docx.Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs]
    lines = []

    # Add part info if this is the first chapter of a part
    if part_info:
        lines.append(f"> **{part_info}**")
        lines.append("")

    lines.append(f"# {title}")
    lines.append("")

    # Skip the first few paragraphs that are just the chapter title/part header
    skip_titles = True
    i = 0
    while i < len(paragraphs):
        text = paragraphs[i].strip()

        # Skip leading title paragraphs (chapter number, title, part header)
        if skip_titles:
            if not text or text.startswith("第") or text == title or (part_info and any(k in text for k in ["與神一同建造", "跟隨聖靈引導", "被聖靈充滿", "結論"])):
                i += 1
                continue
            skip_titles = False

        # Check for section headings
        prev_empty = (i == 0) or (not paragraphs[i - 1].strip())
        next_empty = (i == len(paragraphs) - 1) or (not paragraphs[i + 1].strip() if i + 1 < len(paragraphs) else True)

        if not text:
            lines.append("")
            i += 1
            continue

        # Detect "禱告提示" sections
        if "禱告" in text and ("提示" in text or "引導" in text) and len(text) < 20:
            lines.append(f"### 🙏 {text}")
            lines.append("")
            i += 1
            continue

        # Detect section titles
        if is_zh_section_title(text, prev_empty, next_empty):
            lines.append(f"## {text}")
            lines.append("")
            i += 1
            continue

        # Regular paragraph
        lines.append(text)
        lines.append("")
        i += 1

    # Clean up multiple blank lines
    result = "\n".join(lines)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip() + "\n"


def convert_en_chapter(paragraphs, start_idx, end_idx, title, part_info, sections):
    """Convert a range of English paragraphs to Markdown."""
    lines = []

    if part_info:
        lines.append(f"> **{part_info}**")
        lines.append("")

    lines.append(f"# {title}")
    lines.append("")

    section_set = set(sections) if sections else set()

    for i in range(start_idx, end_idx):
        text = paragraphs[i].strip()

        if not text:
            lines.append("")
            continue

        # Skip running headers and copyright notices
        if text in EN_RUNNING_HEADERS:
            continue
        if any(skip in text for skip in EN_SKIP_PATTERNS):
            continue

        # Skip "Ibid." standalone references
        if text == "Ibid.":
            lines.append(f"*{text}*")
            lines.append("")
            continue

        # Detect section headings
        if text in section_set:
            if text == "Prayer Prompt" or text == "Prayer Prompts":
                lines.append(f"### 🙏 {text}")
            else:
                lines.append(f"## {text}")
            lines.append("")
            continue

        # Detect standalone scripture references
        if EN_SCRIPTURE_PATTERN.match(text) and len(text) < 80:
            lines.append(f"> *{text}*")
            lines.append("")
            continue

        # Regular paragraph
        lines.append(text)
        lines.append("")

    result = "\n".join(lines)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip() + "\n"


def main():
    os.makedirs(ZH_DIR, exist_ok=True)
    os.makedirs(EN_DIR, exist_ok=True)

    # ── Convert Chinese files ──
    print("=== Converting Chinese files ===")
    for filename, file_id, title, part_info in ZH_FILES:
        filepath = os.path.join(DOCS_DIR, filename)
        if not os.path.exists(filepath):
            print(f"  SKIP (not found): {filename}")
            continue
        md_content = convert_zh_file(filepath, title, part_info)
        out_path = os.path.join(ZH_DIR, f"{file_id}.md")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"  ✓ {out_path}")

    # ── Convert English file ──
    print("\n=== Converting English file ===")
    en_filepath = os.path.join(DOCS_DIR, "Dahl_RaisingSpiritLedKids_ET_wo.docx")
    doc = docx.Document(en_filepath)
    all_paragraphs = [p.text for p in doc.paragraphs]
    total = len(all_paragraphs)

    for idx, chapter_info in enumerate(EN_CHAPTERS):
        start_idx, file_id, title, part_info, sections = chapter_info
        # End index is the start of the next chapter, or end of file
        if idx + 1 < len(EN_CHAPTERS):
            end_idx = EN_CHAPTERS[idx + 1][0]
        else:
            end_idx = total

        md_content = convert_en_chapter(all_paragraphs, start_idx, end_idx, title, part_info, sections)
        out_path = os.path.join(EN_DIR, f"{file_id}.md")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"  ✓ {out_path}")

    print(f"\nDone! Generated {len(ZH_FILES)} Chinese + {len(EN_CHAPTERS)} English Markdown files.")


if __name__ == "__main__":
    main()
